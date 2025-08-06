"""Deposition preparation utilities."""

from __future__ import annotations

import json
from datetime import datetime
from typing import List, Dict, Optional

from openai import OpenAI
from docx import Document as DocxDocument
from weasyprint import HTML

from apps.legal_discovery.database import db
from apps.legal_discovery.models import (
    Witness,
    Fact,
    Document,
    DepositionQuestion,
    FactConflict,
    Agent,
    DepositionReviewLog,
)

PROMPT_TMPL = (
    "You are preparing for a deposition of {name}, who is mentioned in:\n{facts}\n"
    "Generate a JSON array of questions with fields 'category', 'question', and 'source' "
    "grouped into Background, Events, Inconsistencies, and Damages."
)


class DepositionPrep:
    """Utilities for deposition preparation."""

    @staticmethod
    def generate_questions(
        witness_id: int, scope: Optional[Dict] = None, include_privileged: bool = False
    ) -> List[Dict]:
        witness = Witness.query.get(witness_id)
        if not witness:
            raise ValueError("Witness not found")

        query = Fact.query.join(Document).filter(Fact.witness_id == witness_id)
        if not include_privileged:
            query = query.filter(Document.is_privileged.is_(False))
        facts = query.all()
        facts_text = "\n".join(
            f"- {f.text} (Doc: {f.document.name})" for f in facts
        ) or "No facts available."

        # Detect contradictions among gathered facts and append to prompt context
        conflicts = DepositionPrep.detect_contradictions(facts, witness_id)
        if conflicts:
            conflict_lines = "\n".join(
                f"- {c['fact1']} <> {c['fact2']}" for c in conflicts
            )
            facts_text += "\nPotential contradictions:\n" + conflict_lines

        prompt = PROMPT_TMPL.format(name=witness.name, facts=facts_text)

        client = OpenAI()
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        content = response.choices[0].message.content

        try:
            data = json.loads(content)
        except json.JSONDecodeError as exc:
            raise ValueError("LLM response not valid JSON") from exc

        DepositionQuestion.query.filter_by(witness_id=witness_id).delete()
        questions = []
        for item in data:
            qobj = DepositionQuestion(
                witness_id=witness_id,
                category=item.get("category", "Misc"),
                question=item.get("question", ""),
                source=item.get("source"),
            )
            db.session.add(qobj)
            questions.append(qobj)
        db.session.commit()

        return [
            {
                "id": q.id,
                "category": q.category,
                "question": q.question,
                "source": q.source,
                "flagged": q.flagged,
            }
            for q in questions
        ]

    @staticmethod
    def detect_contradictions(
        facts: List[Fact], witness_id: int, threshold: float = 0.8
    ) -> List[Dict]:
        """Identify contradictions among witness facts using an LLM."""

        conflicts: List[Dict] = []
        client = OpenAI()
        for i in range(len(facts)):
            for j in range(i + 1, len(facts)):
                prompt = (
                    "Do these statements contradict each other?\n"
                    f"1. {facts[i].text}\n2. {facts[j].text}\n"
                    "Respond with JSON {\"contradiction\": bool, \"score\": float}."
                )
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0,
                )
                content = response.choices[0].message.content
                try:
                    result = json.loads(content)
                except json.JSONDecodeError:
                    continue
                if result.get("contradiction") and result.get("score", 0) >= threshold:
                    conflict = FactConflict(
                        witness_id=witness_id,
                        fact1_id=facts[i].id,
                        fact2_id=facts[j].id,
                        score=float(result.get("score", 0)),
                        description=f'"{facts[i].text}" <> "{facts[j].text}"',
                    )
                    db.session.add(conflict)
                    conflicts.append(
                        {
                            "fact1": facts[i].text,
                            "fact2": facts[j].text,
                            "score": float(result.get("score", 0)),
                        }
                    )
        if conflicts:
            db.session.commit()
        return conflicts

    @staticmethod
    def log_review(
        witness_id: int,
        reviewer_id: int,
        approved: bool,
        notes: Optional[str] = None,
    ) -> Dict:
        reviewer = Agent.query.get(reviewer_id)
        if not reviewer or reviewer.role not in {"attorney", "case_admin"}:
            raise PermissionError("Reviewer lacks permission")
        witness = Witness.query.get_or_404(witness_id)
        log = DepositionReviewLog(
            reviewer_id=reviewer_id,
            witness_id=witness.id,
            approved=approved,
            notes=notes,
        )
        db.session.add(log)
        db.session.commit()
        return {
            "id": log.id,
            "approved": log.approved,
            "notes": log.notes,
        }

    @staticmethod
    def export_questions(
        witness_id: int, file_path: str, reviewer_id: int
    ) -> str:
        """Export deposition questions to PDF or DOCX for an authorized reviewer.

        Returns:
            str: Path to the generated document.
        """

        reviewer = Agent.query.get(reviewer_id)
        if not reviewer or reviewer.role not in {"attorney", "case_admin"}:
            raise PermissionError("Reviewer lacks permission")

        final_path = str(file_path)
        witness = Witness.query.get_or_404(witness_id)
        questions = (
            DepositionQuestion.query.filter_by(witness_id=witness_id)
            .order_by(DepositionQuestion.id)
            .all()
        )
        case_id = witness.associated_case
        timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

        if final_path.lower().endswith(".pdf"):
            items_html = ""
            sources_html = ""
            for idx, q in enumerate(questions, 1):
                items_html += f"<li>{q.question}"
                if q.source:
                    items_html += f"<sup><a href='#src{idx}'>{idx}</a></sup>"
                    sources_html += (
                        f"<li id='src{idx}'><a href='{q.source}'>{q.source}</a></li>"
                    )
                items_html += "</li>"
            html = f"""
            <h1>Deposition Outline: {witness.name}</h1>
            <p>Case ID: {case_id}</p>
            <p>Generated: {timestamp}</p>
            <ol>{items_html}</ol>
            <h2>Sources</h2>
            <ol>{sources_html}</ol>
            """
            HTML(string=html).write_pdf(final_path)
        else:
            doc = DocxDocument()
            doc.add_heading(f"Deposition Outline: {witness.name}", level=1)
            doc.add_paragraph(f"Case ID: {case_id}")
            doc.add_paragraph(f"Generated: {timestamp}")
            sources: List[str] = []
            for q in questions:
                p = doc.add_paragraph(style="List Number")
                p.add_run(q.question)
                if q.source:
                    ref = len(sources) + 1
                    p.add_run(f" [{ref}]")
                    sources.append(q.source)
            if sources:
                doc.add_heading("Sources", level=2)
                for i, src in enumerate(sources, 1):
                    doc.add_paragraph(f"[{i}] {src}")
            doc.save(final_path)

        return final_path

    @staticmethod
    def flag_question(question_id: int) -> None:
        question = DepositionQuestion.query.get_or_404(question_id)
        question.flagged = True
        db.session.commit()
