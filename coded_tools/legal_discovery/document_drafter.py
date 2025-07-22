from docx import Document
from neuro_san.coded_tool import CodedTool


class DocumentDrafter(CodedTool):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

    def create_document(self, filepath: str, content: str):
        """
        Creates a new Word document.

        :param filepath: The path to the new document.
        :param content: The content to add to the document.
        """
        document = Document()
        document.add_paragraph(content)
        document.save(filepath)

    def add_paragraph(self, filepath: str, content: str):
        """
        Adds a new paragraph to an existing Word document.

        :param filepath: The path to the document.
        :param content: The content to add.
        """
        document = Document(filepath)
        document.add_paragraph(content)
        document.save(filepath)

    def add_heading(self, filepath: str, text: str, level: int):
        """
        Adds a new heading to an existing Word document.

        :param filepath: The path to the document.
        :param text: The text of the heading.
        :param level: The level of the heading (1-9).
        """
        document = Document(filepath)
        document.add_heading(text, level=level)
        document.save(filepath)
